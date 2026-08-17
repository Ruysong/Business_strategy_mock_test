from pathlib import Path
import re

import fitz
from docx import Document
from PIL import Image
from rapidocr_onnxruntime import RapidOCR

SOURCE = Path(r"C:\성균관대\10_학교\JEAN\경전 족보\경전 족보")
OUTPUT = Path("ocr-md")
OUTPUT.mkdir(exist_ok=True)
ocr = RapidOCR()
LATEST_STEMS = {"23-2", "경전 23-2 (오프) 족보", "경전 24-1 기말", "경전 25-1 기말"}


def tidy(text):
    lines = [re.sub(r"\s+", " ", line).strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line)


def recognize(pix):
    if isinstance(pix, Image.Image):
        image = pix.convert("RGB")
    else:
        mode = "RGBA" if pix.alpha else "RGB"
        image = Image.frombytes(mode, [pix.width, pix.height], pix.samples)
    result, _ = ocr(image)
    if not result:
        return ""
    rows = sorted(result, key=lambda row: (min(p[1] for p in row[0]), min(p[0] for p in row[0])))
    return "\n".join(row[1] for row in rows)


def pdf_markdown(path):
    doc = fitz.open(path)
    sections = [f"# {path.stem}", "", f"- 원본: `{path.name}`", f"- 총 페이지: {len(doc)}", ""]
    for index, page in enumerate(doc, 1):
        embedded = tidy(page.get_text("text"))
        if len(embedded) >= 120:
            text, method = embedded, "내장 텍스트"
        else:
            pix = page.get_pixmap(matrix=fitz.Matrix(2.2, 2.2), alpha=False)
            text, method = tidy(recognize(pix)), "OCR"
        sections += [f"## 페이지 {index}", "", f"<!-- 추출 방식: {method} -->", "", text or "[텍스트를 인식하지 못함]", ""]
        print(f"{path.name}: {index}/{len(doc)} {method} {len(text)} chars", flush=True)
    return "\n".join(sections)


def image_markdown(path):
    image = Image.open(path).convert("RGB")
    text = tidy(recognize(image))
    return f"# {path.stem}\n\n- 원본: `{path.name}`\n- 총 페이지: 1\n\n## 페이지 1\n\n<!-- 추출 방식: OCR -->\n\n{text or '[텍스트를 인식하지 못함]'}\n"


def docx_markdown(path):
    doc = Document(path)
    chunks = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            chunks.append(" | ".join(cell.text for cell in row.cells))
    text = tidy("\n".join(chunks))
    # Embedded images are also OCRed so image-only DOCX content is not skipped.
    image_texts = []
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                image = Image.open(rel.target_part.blob).convert("RGB")
                recognized = tidy(recognize(image))
                if recognized:
                    image_texts.append(recognized)
            except Exception:
                continue
    sections = [f"# {path.stem}", "", f"- 원본: `{path.name}`", "", "## 문서 텍스트", "", text or "[텍스트 없음]", ""]
    for index, recognized in enumerate(image_texts, 1):
        sections += [f"## 포함 이미지 {index} OCR", "", recognized, ""]
    return "\n".join(sections)


for path in sorted(SOURCE.iterdir()):
    if path.stem not in LATEST_STEMS:
        continue
    target = OUTPUT / f"{path.stem}.md"
    if target.exists():
        print(f"SKIP {target}", flush=True)
        continue
    if path.suffix.lower() == ".pdf":
        content = pdf_markdown(path)
    elif path.suffix.lower() in {".jpg", ".jpeg", ".png"}:
        content = image_markdown(path)
    elif path.suffix.lower() == ".docx":
        content = docx_markdown(path)
    else:
        continue
    target.write_text(content, encoding="utf-8")
    print(f"WROTE {target}", flush=True)
