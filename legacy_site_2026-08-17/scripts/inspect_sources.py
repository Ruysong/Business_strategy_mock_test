from pathlib import Path
import json
import sys

SOURCE = Path(r"C:\성균관대\10_학교\JEAN\경전 족보\경전 족보")


def pdf_info(path):
    import fitz

    doc = fitz.open(path)
    pages = []
    for page in doc:
        text = page.get_text("text").strip()
        pages.append({"chars": len(text), "sample": text[:300]})
    return {
        "pages": len(doc),
        "text_chars": sum(page["chars"] for page in pages),
        "page_samples": pages,
    }


def docx_info(path):
    from docx import Document

    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    table_text = "\n".join(
        cell.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
        if cell.text.strip()
    )
    return {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "text_chars": len(text) + len(table_text),
        "sample": (text + "\n" + table_text)[:1000],
    }


def main():
    result = []
    for path in sorted(SOURCE.iterdir()):
        item = {"name": path.name, "suffix": path.suffix.lower(), "size": path.stat().st_size}
        try:
            if path.suffix.lower() == ".pdf":
                item.update(pdf_info(path))
            elif path.suffix.lower() == ".docx":
                item.update(docx_info(path))
        except Exception as exc:
            item["error"] = f"{type(exc).__name__}: {exc}"
        result.append(item)
    Path("source-inspection.json").write_text(
        json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8"
    )


if __name__ == "__main__":
    main()
