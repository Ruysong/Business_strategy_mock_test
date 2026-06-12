from pathlib import Path
import fitz
from docx import Document

SOURCE = Path(r"C:\성균관대\10_학교\JEAN\경전 족보\경전 족보")
OUTPUT = Path("extracted")
OUTPUT.mkdir(exist_ok=True)

for path in sorted(SOURCE.iterdir()):
    target = OUTPUT / f"{path.stem}.txt"
    if path.suffix.lower() == ".pdf":
        doc = fitz.open(path)
        chunks = []
        for index, page in enumerate(doc, 1):
            chunks.append(f"\n\n===== PAGE {index} =====\n\n{page.get_text('text')}")
        target.write_text("".join(chunks), encoding="utf-8")
    elif path.suffix.lower() == ".docx":
        doc = Document(path)
        chunks = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                chunks.append(" | ".join(cell.text for cell in row.cells))
        target.write_text("\n".join(chunks), encoding="utf-8")
